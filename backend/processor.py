"""
NexusDocs Document Processor Module
Handles text extraction from various file formats and text chunking.
"""

import os
from pathlib import Path
from typing import List, Optional
from dataclasses import dataclass
import hashlib

from pypdf import PdfReader
from docx import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import CHUNK_SIZE, CHUNK_OVERLAP, SUPPORTED_EXTENSIONS


@dataclass
class DocumentChunk:
    """Represents a chunk of text from a document."""
    content: str
    metadata: dict
    chunk_index: int


@dataclass
class ProcessedDocument:
    """Represents a fully processed document."""
    doc_id: str
    filename: str
    folder: str
    file_type: str
    total_chunks: int
    chunks: List[DocumentChunk]


class DocumentProcessor:
    """Handles document extraction and chunking logic."""
    
    def __init__(
        self,
        chunk_size: int = CHUNK_SIZE,
        chunk_overlap: int = CHUNK_OVERLAP
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
    
    def generate_doc_id(self, filename: str, folder: str) -> str:
        """Generate a unique document ID based on filename and folder."""
        unique_string = f"{folder}/{filename}"
        return hashlib.md5(unique_string.encode()).hexdigest()[:12]
    
    def extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from a PDF file."""
        try:
            reader = PdfReader(str(file_path))
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text() or ""
                if page_text.strip():
                    text_parts.append(f"[Page {page_num}]\n{page_text}")
            return "\n\n".join(text_parts)
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from a DOCX file."""
        try:
            doc = Document(str(file_path))
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    def extract_text_from_txt(self, file_path: Path) -> str:
        """Extract text from a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            raise ValueError(f"Failed to extract text from TXT: {e}")
    
    def extract_text_from_md(self, file_path: Path) -> str:
        """Extract text from a Markdown file."""
        # For MD files, we keep the markdown formatting as it provides context
        return self.extract_text_from_txt(file_path)
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from a file based on its extension."""
        extension = file_path.suffix.lower()
        
        if extension not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")
        
        extractors = {
            '.pdf': self.extract_text_from_pdf,
            '.docx': self.extract_text_from_docx,
            '.txt': self.extract_text_from_txt,
            '.md': self.extract_text_from_md,
        }
        
        return extractors[extension](file_path)
    
    def chunk_text(
        self,
        text: str,
        doc_id: str,
        filename: str,
        folder: str,
        file_type: str
    ) -> List[DocumentChunk]:
        """Split text into chunks with metadata."""
        chunks = self.text_splitter.split_text(text)
        
        return [
            DocumentChunk(
                content=chunk,
                metadata={
                    "doc_id": doc_id,
                    "filename": filename,
                    "folder": folder,
                    "file_type": file_type,
                    "chunk_index": i,
                    "total_chunks": len(chunks)
                },
                chunk_index=i
            )
            for i, chunk in enumerate(chunks)
        ]
    
    def process_document(
        self,
        file_path: Path,
        folder: str = "default"
    ) -> ProcessedDocument:
        """
        Process a document: extract text and split into chunks.
        
        Args:
            file_path: Path to the document file
            folder: Folder name for organization
            
        Returns:
            ProcessedDocument with chunks and metadata
        """
        filename = file_path.name
        file_type = file_path.suffix.lower()
        doc_id = self.generate_doc_id(filename, folder)
        
        # Extract text
        text = self.extract_text(file_path)
        
        if not text.strip():
            raise ValueError(f"No text content extracted from {filename}")
        
        # Create chunks
        chunks = self.chunk_text(text, doc_id, filename, folder, file_type)
        
        return ProcessedDocument(
            doc_id=doc_id,
            filename=filename,
            folder=folder,
            file_type=file_type,
            total_chunks=len(chunks),
            chunks=chunks
        )


# Singleton instance for reuse
processor = DocumentProcessor()
